import os
from datetime import datetime

from docx import Document


BASE_FOLDER = "generated"
FINANCE_FOLDER = os.path.join(BASE_FOLDER, "finance")


def create_finance_folder(report_name):
    safe_name = report_name.replace(" ", "_").replace("/", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    folder = os.path.join(
        FINANCE_FOLDER,
        f"{timestamp}_{safe_name}"
    )

    os.makedirs(folder, exist_ok=True)

    return folder


def create_docx(path, title, sections):
    document = Document()

    document.add_heading(title, 0)

    for heading, content in sections:

        document.add_heading(heading, level=1)

        if isinstance(content, list):

            for item in content:
                document.add_paragraph(str(item), style="List Bullet")

        else:
            document.add_paragraph(str(content))

    document.save(path)


def export_financial_summary(folder, report):

    summary = report.get("financial_summary", {})

    sections = [
        ("Monthly Revenue", summary.get("monthly_revenue", "")),
        ("Monthly Expenses", summary.get("monthly_expenses", "")),
        ("Estimated Profit", summary.get("estimated_profit", "")),
        ("Profit Margin", summary.get("profit_margin", "")),
        ("Financial Health", summary.get("financial_health", "")),
    ]

    path = os.path.join(folder, "financial_summary.docx")

    create_docx(path, "Financial Summary", sections)

    return path


def export_expense_analysis(folder, report):

    sections = []

    for item in report.get("expense_analysis", []):

        sections.append((
            item.get("category", ""),
            f"""
Observation:
{item.get('observation','')}

Recommendation:
{item.get('recommendation','')}
"""
        ))

    path = os.path.join(folder, "expense_analysis.docx")

    create_docx(path, "Expense Analysis", sections)

    return path


def export_budget(folder, report):

    sections = []

    for item in report.get("budget_plan", []):

        sections.append((
            item.get("department", ""),
            f"""
Suggested Budget:
{item.get('suggested_budget','')}

Reason:
{item.get('reason','')}
"""
        ))

    path = os.path.join(folder, "budget_plan.docx")

    create_docx(path, "Budget Plan", sections)

    return path


def export_profitability(folder, report):

    profitability = report.get("profitability_report", {})

    sections = [
        (
            "Revenue Drivers",
            profitability.get("revenue_drivers", [])
        ),
        (
            "Cost Drivers",
            profitability.get("cost_drivers", [])
        ),
        (
            "Profit Improvement Actions",
            profitability.get("profit_improvement_actions", [])
        ),
    ]

    path = os.path.join(folder, "profitability_report.docx")

    create_docx(path, "Profitability Report", sections)

    return path


def export_forecast(folder, report):

    forecast = report.get("forecast", {})

    sections = [
        ("Next 30 Days", forecast.get("next_30_days", "")),
        ("Next 90 Days", forecast.get("next_90_days", "")),
        ("Next 12 Months", forecast.get("next_12_months", "")),
    ]

    path = os.path.join(folder, "financial_forecast.docx")

    create_docx(path, "Financial Forecast", sections)

    return path


def export_investments(folder, report):

    sections = []

    for item in report.get("investment_recommendations", []):

        sections.append((
            item.get("recommendation", ""),
            f"""
Priority:
{item.get('priority','')}

Reason:
{item.get('reason','')}
"""
        ))

    path = os.path.join(folder, "investment_recommendations.docx")

    create_docx(path, "Investment Recommendations", sections)

    return path


def export_finance_deliverables(report):

    report_name = report.get("report_name", "Finance Report")

    folder = create_finance_folder(report_name)

    files = {
        "financial_summary_docx":
            export_financial_summary(folder, report),

        "expense_analysis_docx":
            export_expense_analysis(folder, report),

        "budget_plan_docx":
            export_budget(folder, report),

        "profitability_report_docx":
            export_profitability(folder, report),

        "forecast_docx":
            export_forecast(folder, report),

        "investment_recommendations_docx":
            export_investments(folder, report),
    }

    return {
        "folder_path": folder,
        "files": files
    }