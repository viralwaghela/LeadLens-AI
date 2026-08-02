import os
from datetime import datetime

from docx import Document


BASE_FOLDER = "generated"
HR_FOLDER = os.path.join(BASE_FOLDER, "hr")


def create_hr_folder(package_name):
    safe_name = package_name.replace(" ", "_").replace("/", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    folder = os.path.join(
        HR_FOLDER,
        f"{timestamp}_{safe_name}"
    )

    os.makedirs(folder, exist_ok=True)

    return folder


def create_docx(path, title, sections):
    document = Document()

    document.add_heading(title, 0)

    for heading, content in sections:

        document.add_heading(heading, level=1)

        if isinstance(content, list):

            for item in content:
                document.add_paragraph(str(item), style="List Bullet")

        else:
            document.add_paragraph(str(content))

    document.save(path)


def export_job_description(folder, package):

    jd = package.get("job_description", {})

    sections = [
        ("Title", jd.get("title", "")),
        ("Department", jd.get("department", "")),
        ("Summary", jd.get("summary", "")),
        ("Responsibilities", jd.get("responsibilities", [])),
        ("Requirements", jd.get("requirements", [])),
        ("Salary Range", jd.get("salary_range", "")),
    ]

    path = os.path.join(folder, "job_description.docx")

    create_docx(path, "Job Description", sections)

    return path


def export_candidate_profile(folder, package):

    profile = package.get("candidate_profile", {})

    sections = [
        ("Experience", profile.get("experience", "")),
        ("Education", profile.get("education", "")),
        ("Technical Skills", profile.get("technical_skills", [])),
        ("Soft Skills", profile.get("soft_skills", [])),
    ]

    path = os.path.join(folder, "candidate_profile.docx")

    create_docx(path, "Candidate Profile", sections)

    return path


def export_interview_questions(folder, package):

    questions = package.get("interview_questions", [])

    sections = []

    for item in questions:

        sections.append((
            item.get("type", "Question"),
            item.get("question", "")
        ))

    path = os.path.join(folder, "interview_questions.docx")

    create_docx(path, "Interview Questions", sections)

    return path


def export_scorecard(folder, package):

    scorecard = package.get("evaluation_scorecard", [])

    sections = []

    for item in scorecard:

        sections.append((
            item.get("criteria", ""),
            f"Weight: {item.get('weight', '')}"
        ))

    path = os.path.join(folder, "candidate_scorecard.docx")

    create_docx(path, "Candidate Evaluation Scorecard", sections)

    return path


def export_onboarding(folder, package):

    onboarding = package.get("onboarding_plan", {})

    sections = [
        ("Week 1", onboarding.get("week1", [])),
        ("Week 2", onboarding.get("week2", [])),
        ("Week 3", onboarding.get("week3", [])),
        ("Week 4", onboarding.get("week4", [])),
    ]

    path = os.path.join(folder, "onboarding_plan.docx")

    create_docx(path, "Onboarding Plan", sections)

    return path


def export_performance_review(folder, package):

    review = package.get("performance_review", {})

    sections = [
        ("KPIs", review.get("kpis", [])),
        ("Strengths", review.get("strengths", [])),
        ("Improvement Areas", review.get("improvement_areas", [])),
        ("Goals", review.get("goals", [])),
    ]

    path = os.path.join(folder, "performance_review.docx")

    create_docx(path, "Performance Review", sections)

    return path


def export_hiring_recommendation(folder, package):

    hiring = package.get("hiring_recommendation", {})

    sections = [
        ("Decision", hiring.get("decision", "")),
        ("Reason", hiring.get("reason", "")),
    ]

    path = os.path.join(folder, "hiring_recommendation.docx")

    create_docx(path, "Hiring Recommendation", sections)

    return path


def export_hr_deliverables(package):

    package_name = package.get("package_name", "HR Package")

    folder = create_hr_folder(package_name)

    files = {
        "job_description_docx":
            export_job_description(folder, package),

        "candidate_profile_docx":
            export_candidate_profile(folder, package),

        "interview_questions_docx":
            export_interview_questions(folder, package),

        "scorecard_docx":
            export_scorecard(folder, package),

        "onboarding_docx":
            export_onboarding(folder, package),

        "performance_review_docx":
            export_performance_review(folder, package),

        "hiring_recommendation_docx":
            export_hiring_recommendation(folder, package),
    }

    return {
        "folder_path": folder,
        "files": files
    }