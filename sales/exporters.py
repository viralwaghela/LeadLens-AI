import os
from datetime import datetime

from docx import Document


BASE_FOLDER = "generated"
SALES_FOLDER = os.path.join(BASE_FOLDER, "sales")


def create_sales_folder(campaign_name):
    safe_name = campaign_name.replace(" ", "_").replace("/", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    folder_path = os.path.join(
        SALES_FOLDER,
        f"{timestamp}_{safe_name}"
    )

    os.makedirs(folder_path, exist_ok=True)

    return folder_path


def create_docx(file_path, title, sections):
    document = Document()
    document.add_heading(title, 0)

    for section_title, section_content in sections:
        document.add_heading(section_title, level=1)

        if isinstance(section_content, list):
            for item in section_content:
                document.add_paragraph(str(item), style="List Bullet")
        else:
            document.add_paragraph(str(section_content))

    document.save(file_path)


def export_lead_strategy(folder_path, campaign):
    strategy = campaign.get("lead_strategy", {})

    sections = [
        ("Objective", strategy.get("objective", "")),
        ("Ideal Customer Profile", strategy.get("ideal_customer_profile", "")),
        ("Target Segments", strategy.get("target_segments", [])),
        ("Decision Makers", strategy.get("decision_makers", [])),
        ("Prospecting Channels", strategy.get("prospecting_channels", [])),
        ("Sales Angle", strategy.get("sales_angle", "")),
        ("Primary Offer", strategy.get("primary_offer", "")),
        ("Primary CTA", strategy.get("primary_cta", "")),
    ]

    file_path = os.path.join(folder_path, "lead_strategy.docx")
    create_docx(file_path, "Lead Generation Strategy", sections)

    return file_path


def export_cold_emails(folder_path, campaign):
    emails = campaign.get("cold_emails", [])

    sections = []

    for index, email in enumerate(emails, start=1):
        sections.append((
            f"Cold Email {index}: {email.get('subject', '')}",
            f"{email.get('body', '')}\n\nCTA: {email.get('cta', '')}"
        ))

    file_path = os.path.join(folder_path, "cold_emails.docx")
    create_docx(file_path, "Cold Email Campaign", sections)

    return file_path


def export_whatsapp_messages(folder_path, campaign):
    messages = campaign.get("whatsapp_messages", [])

    sections = []

    for item in messages:
        sections.append((
            item.get("stage", "WhatsApp Message"),
            item.get("message", "")
        ))

    file_path = os.path.join(folder_path, "whatsapp_messages.docx")
    create_docx(file_path, "WhatsApp Outreach", sections)

    return file_path


def export_sales_call_script(folder_path, campaign):
    script = campaign.get("sales_call_script", {})

    objections = []

    for item in script.get("objection_handling", []):
        objections.append(
            f"Objection: {item.get('objection', '')}\nResponse: {item.get('response', '')}"
        )

    sections = [
        ("Opening", script.get("opening", "")),
        ("Discovery Questions", script.get("discovery_questions", [])),
        ("Objection Handling", objections),
        ("Closing Script", script.get("closing_script", "")),
    ]

    file_path = os.path.join(folder_path, "sales_call_script.docx")
    create_docx(file_path, "Sales Call Script", sections)

    return file_path


def export_proposal(folder_path, campaign):
    proposal = campaign.get("proposal", {})

    sections = [
        ("Problem Statement", proposal.get("problem_statement", "")),
        ("Solution", proposal.get("solution", "")),
        ("Deliverables", proposal.get("deliverables", [])),
        ("Pricing Suggestion", proposal.get("pricing_suggestion", "")),
        ("Next Steps", proposal.get("next_steps", "")),
    ]

    file_path = os.path.join(folder_path, "proposal.docx")
    create_docx(file_path, proposal.get("title", "Sales Proposal"), sections)

    return file_path


def export_follow_up_sequence(folder_path, campaign):
    sequence = campaign.get("follow_up_sequence", [])

    sections = []

    for item in sequence:
        sections.append((
            f"{item.get('day', '')} - {item.get('channel', '')}",
            item.get("message", "")
        ))

    file_path = os.path.join(folder_path, "follow_up_sequence.docx")
    create_docx(file_path, "Follow-up Sequence", sections)

    return file_path


def export_prospect_list(folder_path, campaign):
    prospects = campaign.get("prospect_list", [])

    sections = []

    for item in prospects:
        sections.append((
            item.get("segment", "Prospect Segment"),
            f"Why Target: {item.get('why_target', '')}\n\nOutreach Angle: {item.get('outreach_angle', '')}"
        ))

    file_path = os.path.join(folder_path, "prospect_list.docx")
    create_docx(file_path, "Prospect List", sections)

    return file_path


def export_sales_deliverables(campaign):
    campaign_name = campaign.get("campaign_name", "Sales Campaign")

    folder_path = create_sales_folder(campaign_name)

    files = {
        "lead_strategy_docx": export_lead_strategy(folder_path, campaign),
        "prospect_list_docx": export_prospect_list(folder_path, campaign),
        "cold_emails_docx": export_cold_emails(folder_path, campaign),
        "whatsapp_messages_docx": export_whatsapp_messages(folder_path, campaign),
        "sales_call_script_docx": export_sales_call_script(folder_path, campaign),
        "proposal_docx": export_proposal(folder_path, campaign),
        "follow_up_sequence_docx": export_follow_up_sequence(folder_path, campaign),
    }

    return {
        "folder_path": folder_path,
        "files": files
    }