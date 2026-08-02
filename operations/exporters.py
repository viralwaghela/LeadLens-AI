import os
from datetime import datetime

from docx import Document


BASE_FOLDER = "generated"
OPERATIONS_FOLDER = os.path.join(BASE_FOLDER, "operations")


def create_operations_folder(package_name):
    safe_name = package_name.replace(" ", "_").replace("/", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    folder = os.path.join(
        OPERATIONS_FOLDER,
        f"{timestamp}_{safe_name}"
    )

    os.makedirs(folder, exist_ok=True)

    return folder


def create_docx(path, title, sections):
    document = Document()

    document.add_heading(title, 0)

    for heading, content in sections:

        document.add_heading(heading, level=1)

        if isinstance(content, list):
            for item in content:
                document.add_paragraph(str(item), style="List Bullet")
        else:
            document.add_paragraph(str(content))

    document.save(path)


def export_daily_plan(folder, package):

    plan = package.get("daily_operations_plan", {})

    sections = [
        ("Objective", plan.get("objective", "")),
        ("Top Priorities", plan.get("top_priorities", [])),
        ("Department Focus", plan.get("department_focus", [])),
        ("Expected Outcome", plan.get("expected_outcome", "")),
    ]

    path = os.path.join(folder, "daily_operations_plan.docx")

    create_docx(path, "Daily Operations Plan", sections)

    return path


def export_task_assignments(folder, package):

    sections = []

    for task in package.get("task_assignment", []):

        sections.append((
            task.get("task", ""),
            f"""
Department:
{task.get('department','')}

Owner:
{task.get('owner_role','')}

Priority:
{task.get('priority','')}

Deadline:
{task.get('deadline','')}

Success Metric:
{task.get('success_metric','')}
"""
        ))

    path = os.path.join(folder, "task_assignments.docx")

    create_docx(path, "Task Assignments", sections)

    return path


def export_bottlenecks(folder, package):

    sections = []

    for item in package.get("bottleneck_detection", []):

        sections.append((
            item.get("bottleneck", ""),
            f"""
Impact:
{item.get('impact','')}

Solution:
{item.get('solution','')}
"""
        ))

    path = os.path.join(folder, "bottlenecks.docx")

    create_docx(path, "Bottleneck Report", sections)

    return path


def export_risks(folder, package):

    sections = []

    for risk in package.get("operational_risks", []):

        sections.append((
            risk.get("risk", ""),
            f"""
Severity:
{risk.get('severity','')}

Mitigation:
{risk.get('mitigation','')}
"""
        ))

    path = os.path.join(folder, "operational_risks.docx")

    create_docx(path, "Operational Risks", sections)

    return path


def export_process_improvements(folder, package):

    sections = []

    for item in package.get("process_improvements", []):

        sections.append((
            item.get("process", ""),
            f"""
Recommendation:
{item.get('recommendation','')}

Expected Benefit:
{item.get('expected_benefit','')}
"""
        ))

    path = os.path.join(folder, "process_improvements.docx")

    create_docx(path, "Process Improvements", sections)

    return path


def export_weekly_report(folder, package):

    report = package.get("weekly_operations_report", {})

    sections = [
        ("Summary", report.get("summary", "")),
        ("Completed Work", report.get("completed_work", [])),
        ("Pending Work", report.get("pending_work", [])),
        ("Next Week Focus", report.get("next_week_focus", [])),
    ]

    path = os.path.join(folder, "weekly_operations_report.docx")

    create_docx(path, "Weekly Operations Report", sections)

    return path


def export_operations_deliverables(package):

    package_name = package.get(
        "package_name",
        "Operations Package"
    )

    folder = create_operations_folder(package_name)

    files = {
        "daily_plan_docx":
            export_daily_plan(folder, package),

        "task_assignments_docx":
            export_task_assignments(folder, package),

        "bottlenecks_docx":
            export_bottlenecks(folder, package),

        "operational_risks_docx":
            export_risks(folder, package),

        "process_improvements_docx":
            export_process_improvements(folder, package),

        "weekly_report_docx":
            export_weekly_report(folder, package),
    }

    return {
        "folder_path": folder,
        "files": files
    }