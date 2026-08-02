import os
from datetime import datetime

from docx import Document
from openpyxl import Workbook


BASE_FOLDER = "generated"
MARKETING_FOLDER = os.path.join(BASE_FOLDER, "marketing")


def create_campaign_folder(campaign_name):
    safe_name = campaign_name.replace(" ", "_").replace("/", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    folder_path = os.path.join(
        MARKETING_FOLDER,
        f"{timestamp}_{safe_name}"
    )

    os.makedirs(folder_path, exist_ok=True)

    return folder_path


def create_docx(file_path, title, sections):
    document = Document()

    document.add_heading(title, 0)

    for section_title, section_content in sections:
        document.add_heading(section_title, level=1)

        if isinstance(section_content, list):
            for item in section_content:
                document.add_paragraph(str(item), style="List Bullet")
        else:
            document.add_paragraph(str(section_content))

    document.save(file_path)


def export_strategy(folder_path, campaign):
    strategy = campaign.get("strategy", {})

    sections = [
        ("Objective", strategy.get("objective", "")),
        ("Target Audience", strategy.get("target_audience", "")),
        ("Pain Points", strategy.get("pain_points", [])),
        ("Positioning", strategy.get("positioning", "")),
        ("Offer", strategy.get("offer", "")),
        ("Primary CTA", strategy.get("primary_cta", "")),
    ]

    file_path = os.path.join(folder_path, "strategy.docx")
    create_docx(file_path, "Marketing Strategy", sections)

    return file_path


def export_reel_scripts(folder_path, campaign):
    reels = campaign.get("reel_ideas", [])

    sections = []

    for reel in reels:
        sections.append((
            reel.get("title", "Untitled Reel"),
            f"Hook: {reel.get('hook', '')}\n\nScript: {reel.get('script', '')}\n\nCTA: {reel.get('cta', '')}"
        ))

    file_path = os.path.join(folder_path, "reel_scripts.docx")
    create_docx(file_path, "Reel Scripts", sections)

    return file_path


def export_captions(folder_path, campaign):
    captions = campaign.get("captions", [])

    sections = []

    for item in captions:
        sections.append((
            item.get("platform", "Platform"),
            item.get("caption", "")
        ))

    file_path = os.path.join(folder_path, "captions.docx")
    create_docx(file_path, "Captions", sections)

    return file_path


def export_meta_ads(folder_path, campaign):
    ads = campaign.get("meta_ads", [])

    sections = []

    for index, ad in enumerate(ads, start=1):
        sections.append((
            f"Meta Ad {index}",
            f"Primary Text: {ad.get('primary_text', '')}\n\nHeadline: {ad.get('headline', '')}\n\nDescription: {ad.get('description', '')}\n\nCTA: {ad.get('cta', '')}"
        ))

    file_path = os.path.join(folder_path, "meta_ads.docx")
    create_docx(file_path, "Meta Ads Copy", sections)

    return file_path


def export_content_calendar(folder_path, campaign):
    calendar = campaign.get("content_calendar", [])

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Content Calendar"

    sheet.append(["Day", "Platform", "Content Type", "Topic", "Goal"])

    for item in calendar:
        sheet.append([
            item.get("day", ""),
            item.get("platform", ""),
            item.get("content_type", ""),
            item.get("topic", ""),
            item.get("goal", ""),
        ])

    file_path = os.path.join(folder_path, "content_calendar.xlsx")
    workbook.save(file_path)

    return file_path


def export_hashtags(folder_path, campaign):
    hashtags = campaign.get("hashtags", [])

    file_path = os.path.join(folder_path, "hashtags.txt")

    with open(file_path, "w", encoding="utf-8") as file:
        file.write(" ".join(hashtags))

    return file_path


def export_image_prompts(folder_path, campaign):
    prompts = campaign.get("image_prompts", [])

    file_path = os.path.join(folder_path, "image_prompts.txt")

    with open(file_path, "w", encoding="utf-8") as file:
        for item in prompts:
            file.write(f"{item.get('title', '')}\n")
            file.write(f"{item.get('prompt', '')}\n\n")

    return file_path


def export_marketing_deliverables(campaign):
    campaign_name = campaign.get("campaign_name", "Marketing Campaign")

    folder_path = create_campaign_folder(campaign_name)

    files = {
        "strategy_docx": export_strategy(folder_path, campaign),
        "calendar_xlsx": export_content_calendar(folder_path, campaign),
        "reel_scripts_docx": export_reel_scripts(folder_path, campaign),
        "captions_docx": export_captions(folder_path, campaign),
        "hashtags_txt": export_hashtags(folder_path, campaign),
        "image_prompts_txt": export_image_prompts(folder_path, campaign),
        "meta_ads_docx": export_meta_ads(folder_path, campaign),
    }

    return {
        "folder_path": folder_path,
        "files": files
    }